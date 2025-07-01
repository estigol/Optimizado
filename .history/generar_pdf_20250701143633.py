import os
from docx import Document
import win32com.client
import pythoncom  # Necesario para inicializar la COM en hilos de Flask

def generar_pdf(datos):
    """
    Genera un PDF a partir de la plantilla Word usando los datos del paciente.
    Retorna la ruta absoluta al PDF generado.
    """
    # Cargar plantilla
    doc = Document("HOJA DE RUTA PROCESO DE SOLICITUD EVALUACION Y CERTIFICACION.docx")
    doc_id = datos["numero_documento"]

    # Mapeo de marcadores en el docx con los campos del formulario
    mapeo = {
        "«Número»": datos["numero_documento"],
        "«NOMBRES_Y_APELLIDOS»": datos["nombres"] + " " + datos["apellidos"],
        "«DOCUMENTO»": datos["tipo_documento"],
        "«Expedida_en»": datos["ciudad_expedicion_documento"],
        "«NACIONALIDAD»": datos["nacionalidad"],
        "«Fecha_Nacimiento»": datos["fecha_nacimiento"],
        "«GÉNERO»": datos["genero"],
        "«Grupo_Sanguíneo»": datos["grupo_sanguineo"],
        "«Estado_Civil»": datos["estado_civil"],
        "«Dirección»": datos["direccion"],
        "«ciudad»": datos["ciudad"],
        "«Correo_Electrónico_»": datos["correo"],
        "«Teléfono_»": datos["telefono"],
        "«Ocupación»": datos["ocupacion"],
        "«Escolaridad»": datos["escolaridad"],
        "«EPS»": datos["eps"],
        "«TIPO_VINCULACION»": datos["tipo_vinculacion"],
        "«MunicipioCiudad_Residencia»": datos["ciudad"],
        "«En_caso_de_emergencia_avisar_a»": datos["contacto_emergencia"],
        "«Teléfono__3»": datos["telefono_emergencia"],
        "«Parentesco»": datos["parentesco_emergencia"],
        "«Acompañante_Solo_menores_de_edad»": datos["acompanante_menor"],
        "«Moto»": datos["tramite_moto"],
        "«TRAMITEMOTO»": datos["tramite_moto"],
        "«CATEGORIAMOTO»": datos["categoria_moto"],
        "«TRAMITECARRO»": datos["tramite_carro"],
        "«Carro»": datos["tramite_carro"],
        "«CATEGORIACARRO»": datos["categoria_carro"],
        "«A_Licencia_»": datos["tramite_carro"],
        "«Usa_GafasLentes_Correctivos»": datos["usa_lentes"],
        "«Usa_Audífonos»": datos["usa_audifonos"],
        "«Usa_Prótesis»": datos["tiene_protesis"],
        "«Usa_Marcapasos»": datos["tiene_marcapasos"],
        "«Consume_medicamentos_»": datos["consume_medicamentos"],
        "«Cuales»": datos["medicamentos"],
        "«Le_han_hecho_cirugías»": datos["ha_tenido_cirugias"],
        "«Cuales_4»": datos["cirugias"],
        "«Consume_Alcohol»": datos["consume_alcohol"],
        "«Ha_estado_en_tratamiento_psicológicopsi»": datos["tratamiento_psico"]
    }

    # Reemplazo de marcadores en el documento
    def reemplazar_en_doc(doc, marcador, valor):
        for p in doc.paragraphs:
            if marcador in p.text:
                for r in p.runs:
                    r.text = r.text.replace(marcador, str(valor))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    reemplazar_en_doc(cell, marcador, valor)

    for marcador, valor in mapeo.items():
        reemplazar_en_doc(doc, marcador, valor)

    # Guardar el .docx y convertir a PDF
    os.makedirs("PDFsGenerados", exist_ok=True)
    output_docx = os.path.abspath(f"PDFsGenerados/{doc_id}_hoja_ruta.docx")
    output_pdf = os.path.abspath(f"PDFsGenerados/{doc_id}_hoja_ruta.pdf")
    doc.save(output_docx)

    # Inicializar COM y convertir docx a PDF usando Word (requiere Windows)
    pythoncom.CoInitialize()  # <- ¡Esta línea soluciona tu error!
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        word_doc = word.Documents.Open(output_docx)
        word_doc.SaveAs(output_pdf, FileFormat=17)  # 17 = PDF
        word_doc.Close()
    finally:
        word.Quit()

    return output_pdf
